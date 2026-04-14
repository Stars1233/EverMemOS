"""GDPVal evaluator — adapted from ClawWork's LLMEvaluator.

Uses category-specific meta-prompts (per occupation) to evaluate agent
deliverables via an LLM.  The original calls OpenAI directly; this
version routes through OpenRouter so we can reuse the same API key.

Source: https://github.com/HKUDS/ClawWork/blob/main/livebench/work/llm_evaluator.py
Changes:
  - OpenAI client → OpenRouter (base_url + model prefix)
  - Removed livebench.utils.logger dependency (replaced with stdlib logging)
  - Inlined read_pdf_as_images / read_pptx_as_images helpers
"""

import base64
import io
import json
import logging
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union

from openai import OpenAI

log = logging.getLogger("evoagentbench")

# ── File-reading helpers (from ClawWork livebench/tools/productivity/file_reading.py) ─

def _read_pdf_as_images(pdf_path: Path) -> Optional[List[bytes]]:
    """Convert PDF → list of PNG bytes (4 pages per combined 2×2 image)."""
    try:
        from PIL import Image
        from pdf2image import convert_from_path
    except ImportError:
        log.warning("pdf2image or Pillow not installed; falling back to text extraction for PDF")
        return None

    try:
        images = convert_from_path(str(pdf_path), dpi=100)
        if not images:
            return None

        combined_images = []
        for i in range(0, len(images), 4):
            batch = images[i:i + 4]
            resized = []
            for img in batch:
                if img.width > 600:
                    ratio = 600 / img.width
                    img = img.resize((600, int(img.height * ratio)), Image.Resampling.LANCZOS)
                resized.append(img)

            max_w = max(im.width for im in resized)
            max_h = max(im.height for im in resized)
            cols = 2
            rows = (len(resized) + 1) // 2
            combined = Image.new("RGB", (max_w * cols, max_h * rows), "white")
            for idx, img in enumerate(resized):
                combined.paste(img, ((idx % cols) * max_w, (idx // cols) * max_h))

            buf = io.BytesIO()
            combined.save(buf, format="PNG", optimize=True)
            combined_images.append(buf.getvalue())
        return combined_images
    except Exception as e:
        log.warning(f"PDF→image conversion failed: {e}")
        return None


def _read_pptx_as_images(pptx_path: Path) -> Optional[List[bytes]]:
    """Convert PPTX → list of PNG bytes (one per slide) via LibreOffice."""
    try:
        from PIL import Image
        from pdf2image import convert_from_path
    except ImportError:
        return None

    temp_dir = None
    try:
        temp_dir = tempfile.mkdtemp()
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", temp_dir, str(pptx_path)],
            capture_output=True, timeout=30, text=True,
        )
        if result.returncode != 0:
            return None

        pdf_name = pptx_path.stem + ".pdf"
        pdf_path = os.path.join(temp_dir, pdf_name)
        if not os.path.exists(pdf_path):
            return None

        images = convert_from_path(pdf_path, dpi=150)
        out = []
        for img in images:
            if img.width > 1200:
                ratio = 1200 / img.width
                img = img.resize((1200, int(img.height * ratio)), Image.Resampling.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="PNG")
            out.append(buf.getvalue())
        return out
    except Exception:
        return None
    finally:
        if temp_dir:
            shutil.rmtree(temp_dir, ignore_errors=True)


# ── Text-based fallback extractors ────────────────────────────────

def _read_docx_content(path: str) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    doc = Document(path)
    parts = [f"[DOCX - {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables]"]
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for i, table in enumerate(doc.tables):
        parts.append(f"\n--- Table {i+1} ({len(table.rows)}×{len(table.columns)}) ---")
        for row in table.rows[:10]:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
        if len(table.rows) > 10:
            parts.append(f"... ({len(table.rows) - 10} more rows)")
    return "\n".join(parts)


def _read_xlsx_content(path: str) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")
    wb = load_workbook(path, data_only=True)
    parts = [f"[Excel - {len(wb.sheetnames)} sheets: {', '.join(wb.sheetnames)}]"]
    for name in wb.sheetnames[:5]:
        ws = wb[name]
        parts.append(f"\n=== Sheet: {name} ({ws.max_row}×{ws.max_column}) ===")
        for ri, row in enumerate(ws.iter_rows(max_row=20, values_only=True), 1):
            row_text = " | ".join(str(c) if c is not None else "" for c in row)
            if row_text.strip():
                parts.append(f"Row {ri}: {row_text}")
        if (ws.max_row or 0) > 20:
            parts.append(f"... ({ws.max_row - 20} more rows)")
    if len(wb.sheetnames) > 5:
        parts.append(f"\n... ({len(wb.sheetnames) - 5} more sheets)")
    return "\n".join(parts)


# ── Core evaluator ────────────────────────────────────────────────

class LLMEvaluator:
    """Adapted from ClawWork — uses OpenRouter instead of OpenAI directly."""

    def __init__(self, meta_prompts_dir: str, model: str = "openai/gpt-4o",
                 api_key: str | None = None):
        self.meta_prompts_dir = Path(meta_prompts_dir)
        self.model = os.getenv("EVALUATION_MODEL", model)
        api_key = api_key or os.getenv("OPENROUTER_API_KEY", "")
        if not api_key:
            raise ValueError("OPENROUTER_API_KEY not set")

        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=api_key,
        )
        self._cache: Dict[str, Dict] = {}
        log.info(f"LLMEvaluator: model={self.model}")

    # ── meta-prompt loading ──

    def _load_meta_prompt(self, occupation: str) -> Optional[Dict]:
        normalized = occupation.replace(" ", "_").replace(",", "")
        if normalized in self._cache:
            return self._cache[normalized]
        path = self.meta_prompts_dir / f"{normalized}.json"
        if not path.exists():
            log.warning(f"No meta-prompt for occupation '{occupation}' at {path}")
            return None
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        self._cache[normalized] = data
        return data

    # ── artifact reading (multimodal) ──

    def _read_artifacts(self, paths: list[str]) -> Dict[str, Dict]:
        artifacts = {}
        for p in paths:
            size = os.path.getsize(p)
            ext = os.path.splitext(p)[1].lower()

            if size == 0:
                raise ValueError(f"Empty file submitted for evaluation: {p}")
            if size > 2 * 1024 * 1024:
                raise RuntimeError(f"File too large: {size} bytes (>2MB) - {p}")

            if ext in (".png", ".jpg", ".jpeg", ".gif", ".webp"):
                with open(p, "rb") as f:
                    artifacts[p] = {"type": "image", "format": ext[1:], "data": f.read(), "size": size}

            elif ext == ".pdf":
                images = _read_pdf_as_images(Path(p))
                if not images:
                    raise RuntimeError(
                        f"PDF conversion failed for {p}. "
                        f"Ensure poppler-utils and pdf2image are installed.")
                artifacts[p] = {"type": "pdf_images", "images": images,
                                "image_count": len(images), "size": size}

            elif ext == ".pptx":
                images = _read_pptx_as_images(Path(p))
                if not images:
                    raise RuntimeError(
                        f"PPTX conversion failed for {p}. "
                        f"Ensure LibreOffice and pdf2image are installed.")
                artifacts[p] = {"type": "pptx_images", "images": images,
                                "slide_count": len(images), "size": size}

            elif ext == ".docx":
                artifacts[p] = {"type": "text", "content": _read_docx_content(p)}

            elif ext in (".xlsx", ".xls", ".xlsm"):
                artifacts[p] = {"type": "text", "content": _read_xlsx_content(p)}

            else:
                try:
                    with open(p, "r", encoding="utf-8") as f:
                        artifacts[p] = {"type": "text", "content": f.read()}
                except UnicodeDecodeError:
                    raise RuntimeError(f"Unsupported binary file type: {ext} - {p}")

        return artifacts

    # ── build multimodal content for API ──

    def _build_content(self, meta_prompt: Dict, task: Dict,
                       artifact_data: Dict, missing: list[str],
                       description: str) -> List[Dict]:
        evaluation_prompt = meta_prompt.get("evaluation_prompt", "")

        text = f"""# TASK EVALUATION REQUEST

## Category: {meta_prompt.get('category', 'Unknown')}

## Evaluation Guidelines:
{evaluation_prompt}

## Task Prompt (Original Assignment):
{task.get('prompt', 'N/A')}

## Task Metadata:
- Task ID: {task.get('task_id', 'N/A')}
- Sector: {task.get('sector', 'N/A')}
- Occupation: {task.get('occupation', 'N/A')}
- Reference Files: {', '.join(task.get('reference_files', [])) or 'None'}

## Agent's Description:
{description or 'No description provided'}

## Submitted Artifacts:

"""
        for path, art in artifact_data.items():
            name = os.path.basename(path)
            if art["type"] == "text":
                text += f"\n### File: {name}\n```\n{art['content']}\n```\n\n"
            elif art["type"] == "image":
                text += f"\n### Image: {name} ({art['format']}, {art['size']} bytes)\n[See image below]\n\n"
            elif art["type"] == "pptx_images":
                text += f"\n### PowerPoint: {name} ({art['slide_count']} slides)\n[See slide images below]\n\n"
            elif art["type"] == "pdf_images":
                approx = art["image_count"] * 4
                text += f"\n### PDF: {name} (~{approx} pages in {art['image_count']} combined images)\n[See PDF pages below]\n\n"

        if missing:
            text += "\n## Missing Artifacts:\n" + "".join(f"- {p}\n" for p in missing)

        text += """
---

Please evaluate this work according to the rubric above. Output your evaluation in this format:

**OVERALL SCORE:** [0-10]

**DIMENSION SCORES:**
[List dimension scores from rubric]

**KEY FINDINGS:**
[2-3 bullet points on what worked / didn't work]

**FEEDBACK:**
[1-2 paragraph explanation]

**TOP IMPROVEMENTS NEEDED:**
[Numbered list of 3 specific improvements]
"""
        content: List[Dict] = [{"type": "text", "text": text}]

        for path, art in artifact_data.items():
            if art["type"] == "image":
                b64 = base64.b64encode(art["data"]).decode()
                mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                        "gif": "image/gif", "webp": "image/webp"}.get(art["format"], "image/png")
                content.append({"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}", "detail": "high"}})

            elif art["type"] in ("pptx_images", "pdf_images"):
                for img_bytes in art["images"]:
                    b64 = base64.b64encode(img_bytes).decode()
                    content.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}", "detail": "high"}})

        return content

    # ── score extraction ──

    @staticmethod
    def _extract_score(text: str) -> float:
        for pat in [r"OVERALL SCORE:\s*(\d+(?:\.\d+)?)",
                    r"Overall Score:\s*(\d+(?:\.\d+)?)",
                    r"Score:\s*(\d+(?:\.\d+)?)/10",
                    r"Final Score:\s*(\d+(?:\.\d+)?)"]:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                return max(0.0, min(10.0, float(m.group(1))))
        # fallback
        nums = re.findall(r"\b(\d+(?:\.\d+)?)\b", text[:200])
        for n in nums:
            v = float(n)
            if 0 <= v <= 10:
                return v
        log.warning("Could not extract score from evaluation, defaulting to 5.0")
        return 5.0

    # ── main entry point ──

    def evaluate_artifact(self, task: Dict, artifact_paths: list[str],
                          description: str = "",
                          max_payment: float = 50.0) -> Tuple[float, str, float]:
        """Evaluate artifacts against occupation-specific rubric.

        Returns (normalized_score 0-1, feedback_text, payment).
        """
        occupation = task.get("occupation", "")
        if not occupation:
            return 0.0, "Error: no occupation", 0.0

        meta_prompt = self._load_meta_prompt(occupation)
        if not meta_prompt:
            raise FileNotFoundError(f"No meta-prompt for '{occupation}'")

        existing = [p for p in artifact_paths if os.path.exists(p)]
        missing = [p for p in artifact_paths if not os.path.exists(p)]
        if not existing:
            return 0.0, f"No artifacts found: {artifact_paths}", 0.0

        artifact_data = self._read_artifacts(existing)
        content = self._build_content(meta_prompt, task, artifact_data, missing, description)

        try:
            resp = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system",
                     "content": "You are an expert work evaluator. Follow the provided rubric precisely and output a structured evaluation."},
                    {"role": "user", "content": content},
                ],
            )
            eval_text = resp.choices[0].message.content
        except Exception as e:
            raise RuntimeError(f"LLM evaluation failed: {e}") from e

        score = self._extract_score(eval_text)
        normalized = score / 10.0
        payment = normalized * max_payment
        return normalized, eval_text, payment


def _read_pdf_text(path: str) -> str:
    """Fallback: extract text from PDF (when image conversion unavailable)."""
    try:
        import fitz
        doc = fitz.open(path)
        return "\n".join(page.get_text() for page in doc)
    except ImportError:
        pass
    try:
        import pdfplumber
        with pdfplumber.open(path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    except ImportError:
        pass
    return f"[PDF file: {os.path.getsize(path)} bytes, no extraction library available]"


# ── Convenience wrapper for the domain adapter ─────────────────

_ARTIFACT_EXTENSIONS = {
    ".pdf", ".docx", ".xlsx", ".xls", ".xlsm", ".pptx",
    ".txt", ".csv", ".json", ".md", ".py", ".js", ".html", ".css",
    ".png", ".jpg", ".jpeg", ".gif", ".webp",
}

_evaluator: Optional[LLMEvaluator] = None


_MIN_EVALUATION_THRESHOLD = 0.6  # ClawWork payment cliff


def evaluate_rubric(task: dict, workspace_dir: Path, api_key: str,
                    meta_prompts_dir: str = "", model: str = "openai/gpt-4o") -> dict:
    """Evaluate deliverables in workspace using ClawWork's LLMEvaluator.

    Applies the 0.6 payment cliff: scores below 0.6 are set to 0.0
    (consistent with ClawWork/OpenSpace behavior).

    Returns: {"reward": float, "evaluation_score": float, "feedback": str, ...}
    """
    global _evaluator
    if _evaluator is None:
        if not meta_prompts_dir:
            # Default: data/gdpval/meta_prompts relative to project root
            meta_prompts_dir = str(Path(__file__).resolve().parents[3] / "data" / "knowledge_work" / "meta_prompts")
        _evaluator = LLMEvaluator(meta_prompts_dir=meta_prompts_dir, model=model, api_key=api_key)

    # Discover artifacts in workspace (same logic as gdpval_bench)
    artifact_paths = []
    ref_basenames = {os.path.basename(r) for r in task.get("reference_files", [])}
    skip_dirs = {'node_modules', '__pycache__', '.git', 'venv', '.venv'}
    skip_files = {'AGENTS.md', 'BOOTSTRAP.md', 'HEARTBEAT.md', 'IDENTITY.md',
                  'SOUL.md', 'TOOLS.md', 'USER.md', 'MEMORY.md',
                  'workspace-state.json'}
    if workspace_dir.exists():
        for f in sorted(workspace_dir.rglob("*")):
            if not f.is_file():
                continue
            if skip_dirs & set(f.relative_to(workspace_dir).parts):
                continue
            if f.suffix.lower() not in _ARTIFACT_EXTENSIONS:
                continue
            if f.name in ref_basenames or f.name in skip_files:
                continue
            if f.stat().st_size == 0:
                continue
            artifact_paths.append(str(f))

    if not artifact_paths:
        return {"reward": 0.0, "evaluation_score": 0.0,
                "error": "no_artifacts_found", "files_found": []}

    try:
        score, feedback, payment = _evaluator.evaluate_artifact(
            task=task,
            artifact_paths=artifact_paths,
            description=f"Work submission with {len(artifact_paths)} artifact(s)",
            max_payment=task.get("task_value_usd", 50.0),
        )
    except Exception as e:
        log.error(f"  Evaluation failed: {e}")
        return {"reward": 0.0, "evaluation_score": 0.0, "error": str(e),
                "files_found": [os.path.basename(p) for p in artifact_paths]}

    # Apply 0.6 payment cliff (ClawWork-compatible)
    cliff_applied = score < _MIN_EVALUATION_THRESHOLD
    reward = 0.0 if cliff_applied else score

    return {
        "reward": round(reward, 4),
        "evaluation_score": round(score, 4),
        "score_raw": round(score * 10, 2),
        "cliff_applied": cliff_applied,
        "feedback": feedback,
        "payment": round(0.0 if cliff_applied else payment, 2),
        "files_found": [os.path.basename(p) for p in artifact_paths],
    }
